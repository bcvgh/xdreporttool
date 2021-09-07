import docx
from docx import Document
from docx.shared import  Pt
import requests,json,time
import os,re
import pandas as pd
import argparse
from requests.packages.urllib3.exceptions import InsecureRequestWarning
# -*- coding: utf-8 -*-
def arg_s():
    parser = argparse.ArgumentParser(description="deal report")
    parser.add_argument('-u','--update',help='报告更新')
    parser.add_argument('-c','--csv',help='csv报告下载')
    parser.add_argument('-d','--docx',help='csv生成docx')
    parser.add_argument('-r','--replace',help='报告内容替换')
    parser.add_argument('-a', '--all', help='整套流程')
    args = parser.parse_args()
    return args

def mod_t1(file):
    ##高危关键合并
    h_t1 = file.tables[5].cell(1, 2)
    h_t1.text = ''
    ht_t2 = file.tables[5].cell(1, 3)
    n = h_t1.merge(ht_t2)
    n.text = '高危'
    n.paragraphs[0].runs[0].font.size = Pt(12)
    n.paragraphs[0].runs[0].bold = True
    # n.paragraphs[0].runs[0].vertical_alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    h_n1 = file.tables[5].cell(2, 2)
    h_n1_t = int(h_n1.text)
    h_n1.text = ''
    h_n2 = file.tables[5].cell(2, 3)
    h_n2_t = int(h_n2.text)
    h_n2.text = ''
    h1 = h_n1.merge(h_n2)
    sum = str(h_n1_t + h_n2_t)
    h1.text = sum
 ##低危很低危合并
    l_t1 = file.tables[5].cell(1, 5)
    l_t2 = file.tables[5].cell(1, 6)
    l_t2.text = ''
    l_t1.merge(l_t2)
    l_n1 = file.tables[5].cell(2, 5)
    l_n2 = file.tables[5].cell(2, 6)
    l_n2.text = ''
    l_n1.merge(l_n2)

def mod_t2(file):
    row1 = file.tables[-1].rows[2]
    row2 = file.tables[-1].rows[5]
    row1._element.getparent().remove(row1._element)
    row2._element.getparent().remove(row2._element)

def replace_text(file):
    for p in file.paragraphs:
        # 如果要搜索的内容在该段落
        if "关键" in p.text:
            for run in p.runs:
                 if "关键" in run.text:
                    run.text = run.text.replace('关键', '高危')
        if "严重" in p.text:
             for run in p.runs:
                 if "严重" in run.text:
                      run.text = run.text.replace('严重', '高危')

        if "中等" in p.text:
             for run in p.runs:
                 if "中等" in run.text:
                      run.text = run.text.replace('中等', '低危')
    n = 0
    for table in file.tables:
        n = n + 1
        if n > 2:
            for row in table.rows:
                for cell in row.cells:
                    cell.text = cell.text.replace('关键', '高危')
                    cell.text = cell.text.replace('严重', '高危')
                    cell.text = cell.text.replace('中等', '中危')

def mof(f_path,name,out_dir):
    df = pd.read_csv(f_path,header=0)
    pd.set_option('display.max_columns', None)
    pd.set_option('display.max_rows', None)
    pd.set_option('display.width', None)
    try:
        df.rename(columns={'CVSS v2.0 Base Score': 'CVSS'}, inplace=True)
        df.rename(columns={'CVSS v2.0 Temporal Score': 'CVSS Temporal Score'}, inplace=True)
        df.to_csv(out_dir + '\\' + name)
    except Exception:
        return False
    else:
        return True

class Modfied(object):
    def __init__(self,path,out_path,company_name, author,output_path,csv_dir, csv_dir1,password,user,csv_down,out_dir,t_dir):
        burp0_headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:91.0) Gecko/20100101 Firefox/91.0",
                         "Accept": "*/*",
                         "Accept-Language": "zh-CN,zh;q=0.8,zh-TW;q=0.7,zh-HK;q=0.5,en-US;q=0.3,en;q=0.2",
                         "Accept-Encoding": "gzip, deflate",
                         "Content-Type": "application/json",
                         "Origin": "https://192.168.2.113:8834",
                         "Connection": "close",
                         "Referer": "https://192.168.2.113:8834/",
                         "Sec-Fetch-Dest": "empty",
                         "Sec-Fetch-Mode": "cors",
                         "Sec-Fetch-Site": "same-origin",
                         "X-API-Token": "CCA04DCB-6B2D-4F77-9D06-6DB8A121A88B", }
        self.company_name=company_name
        self.author=author
        self.path=path
        self.out_path=out_path
        self.output_path=output_path
        self.csv_dir=csv_dir
        self.csv_dir1=csv_dir1
        self.burp0_url="https://192.168.2.113:8834/session"
        self.burp0_json={"password": password, "username": user}
        self.burp0_headers=burp0_headers
        self.csv_down=csv_down
        self.out_dir=out_dir
        self.t_dir=t_dir

    def erdir(self):
        num = 0
        for i in os.listdir(self.t_dir):
            f_path = self.t_dir + '\\' + i
            res = mof(f_path, i,self.t_dir)
            if res:
                print(i + '文件修改成功!')
                num += 1
            else:
                print(i + '文件修改失败...')
        return num

    def ness_pl(self):
        s = requests.session()
        response = s.post(url=self.burp0_url, json=self.burp0_json, headers=self.burp0_headers, verify=False)
        cookie = 'token=' + json.loads(response.content)["token"]
        self.burp0_headers["X-Cookie"] = cookie

        for n in range(1, 200):
            burp1_url = "https://192.168.2.113:8834/scans/{0}/export?limit=2500".format(n)
            burp0_json = {"extraFilters": {"host_ids": [], "plugin_ids": []}, "format": "csv", "reportContents": {
                "csvColumns": {"cve": True, "cvss": True, "cvss3_base_score": True, "cvss3_temporal_score": True,
                               "cvss_temporal_score": True, "description": True, "exploitable_with": True,
                               "hostname": True, "id": True, "plugin_information": True, "plugin_name": True,
                               "plugin_output": True, "port": True, "protocol": True, "references": True, "risk": True,
                               "risk_factor": True, "see_also": True, "solution": True, "stig_severity": True,
                               "synopsis": True}}}
            res1 = s.post(burp1_url, json=burp0_json, headers=self.burp0_headers, verify=False)
            if 'The requested file was not found' in res1.content.decode():
                continue
            token2 = json.loads(res1.content)["token"]
            burp2_url = "https://192.168.2.113:8834/tokens/{0}/status".format(token2)
            burp3_url = "https://192.168.2.113:8834/tokens/{0}/download".format(token2)
            res2 = s.get(burp2_url, headers=self.burp0_headers, verify=False)
            res3 = s.get(burp3_url, headers=self.burp0_headers, verify=False)
            try:
                pattern = re.compile(r'[0-9]{1,3}_[0-9]{1,3}_[0-9]{1,3}_[0-9]{1,3}')
                ip = re.findall(pattern, res3.headers["Content-Disposition"])
                ip = re.sub('_', '.', ip[0])
            except KeyError:
                print(str(ip) + '号下载失败')
                pass
            else:
                with open(self.csv_down+"\\{0}.csv".format(ip), "wb+") as f:
                    f.write(res3.content)
                print(str(ip) + '下载完成')

            time.sleep(3)

    def update(self):
        for i in os.listdir(self.path):
            f_path = self.path + '\\' + i
            file=Document(f_path)
            mod_t1(file)
            mod_t2(file)
            replace_text(file)
            file.save(self.out_path+'\\'+i)

    def pl_report(self):
        for i in os.listdir(self.csv_dir1):
            pattern = re.compile(r'[0-9]{1,3}.[0-9]{1,3}.[0-9]{1,3}.[0-9]{1,3}')
            ip = re.findall(pattern, i)
            system_name = ip[0]
            csv_file = self.csv_dir + i
            output_path = self.output_path
            com = 'python3 ' + 'NessusReport-v2021042315.py ' + system_name + ' ' + self.company_name + ' ' + self.author + ' ' + csv_file + ' ' + output_path
            os.popen(com)

if __name__ == '__main__':
    res = Modfied(
        path=r'C:\Users\zll\Desktop\rb\aa',
        out_path=r'C:\Users\zll\Desktop\报告\****\漏扫\**',
        company_name="*****",
        author="****",
        output_path='.\\output\\',
        csv_dir="data\\",
        csv_dir1='D:\\bb',
        user='admin',
        password='123456',
        csv_down='D:\\aa\\',
        out_dir='D:\\bb',
        t_dir='D:\\aa'
    )
    args=arg_s()
    if args.all:
        res.ness_pl()
        res.erdir()
        res.pl_report()
        res.update()
    elif args.csv:
        res.ness_pl()
    elif args.update:
        res.erdir()
    elif args.docx:
        res.pl_report()
    elif args.replace:
        res.update()

























# file=Document('D:\\安服脚本\\output\\上海信息网络有限公司_客户管理系统_漏洞扫描报告_v0.1.docx')
# def mod_t1():
#     ##高危关键合并
#     h_t1 = file.tables[5].cell(1, 2)
#     h_t1.text = ''
#     ht_t2 = file.tables[5].cell(1, 3)
#     n = h_t1.merge(ht_t2)
#     n.text = '高危'
#     n.paragraphs[0].runs[0].font.size = Pt(12)
#     n.paragraphs[0].runs[0].bold = True
#     # n.paragraphs[0].runs[0].vertical_alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#     h_n1 = file.tables[5].cell(2, 2)
#     h_n1_t = int(h_n1.text)
#     h_n1.text = ''
#     h_n2 = file.tables[5].cell(2, 3)
#     h_n2_t = int(h_n2.text)
#     h_n2.text = ''
#     h1 = h_n1.merge(h_n2)
#     sum = str(h_n1_t + h_n2_t)
#     h1.text = sum
#
#     ##低危很低危合并
#     l_t1 = file.tables[5].cell(1, 5)
#     l_t2 = file.tables[5].cell(1, 6)
#     l_t2.text = ''
#     l_t1.merge(l_t2)
#     l_n1 = file.tables[5].cell(2, 5)
#     l_n2 = file.tables[5].cell(2, 6)
#     l_n2.text = ''
#     l_n1.merge(l_n2)
#
# def mod_t2():
#     row1=file.tables[-1].rows[2]
#     row2 = file.tables[-1].rows[5]
#     row1._element.getparent().remove(row1._element)
#     row2._element.getparent().remove(row2._element)
#
# def replace_text():
#     for p in file.paragraphs:
#         # 如果要搜索的内容在该段落
#         if "关键" in p.text:
#             for run in p.runs:
#                 if "关键" in run.text:
#                     run.text = run.text.replace('关键','高危')
#         if "严重" in p.text:
#             for run in p.runs:
#                 if "严重" in run.text:
#                     run.text = run.text.replace('严重','高危')
#         if "中等" in p.text:
#             for run in p.runs:
#                 if "中等" in run.text:
#                     run.text = run.text.replace('中等','低危')
#
#
#     n=0
#     for table in file.tables:
#         n=n+1
#         if n>2:
#             for row in table.rows:
#                 for cell in row.cells:
#                     cell.text = cell.text.replace('关键', '高危')
#                     cell.text = cell.text.replace('严重', '高危')
#                     cell.text = cell.text.replace('中等', '中危')
#
# # def mod_style():
# #     ss=file.tables[5].rows[1].cells[3].vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
# #     print(dir(ss))
#
# if __name__ == '__main__':
#
#
